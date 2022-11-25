import docx
picture_1 = "/home/jackchau/Jack/Work/photoWord/Photo/Folder_1/photo1.jpg"
picture_2 = "/home/jackchau/Jack/Work/photoWord/Photo/Folder_1/photo2.jpg"
document = docx.Document ( )

table = document.add_table (rows = 1, cols = 2)

row_cells = table.add_row ( ).cells

left_pic = row_cells[0].paragraphs[0]
right_pic = row_cells[1].paragraphs[0]

run_left = left_pic.add_run ( )
run_right = right_pic.add_run ( )

run_left.add_picture ( picture_1, width=2000000, height = 2000000 )
run_right.add_picture ( picture_2, width=2000000, height = 2000000 )

document.add_table ( rows = 1, cols = 2)

document.save ( "demo.docx" )
