import docx
import os
import numpy as np
from docx.text.run import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import  WD_ORIENT
from docx.shared import Cm, Inches, Pt

doc = docx.Document ( )

set_image_path = "./Photo/"

font_style = "Times New Roman"
font_size  = Pt ( 18 )

image_width = Cm ( 10 )
image_height = Inches ( 3.8 )

# Set margin
margin = Cm ( 2 )
section = doc.sections[ 0 ]
section.top_margin = margin
section.left_margin = margin
section.bottom_margin = margin
section.right_margin = margin

# Connection path
get_folders = os.listdir ( set_image_path )
get_folders.sort ( )
each_folder_path = [ set_image_path + folder for folder in get_folders ]

images_path = [ ]

for path in each_folder_path :
    folder_images = os.listdir ( path )
    folder_images.sort ( )
    each_image_path = [ path + "/" + image for image in folder_images ]
    images_path.append ( each_image_path )

path_arrays = np.array ( images_path, dtype = object )

# Set Orientation
def setOrientation ( ) :
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

# Paste image
def main ( ) :
    for folder in path_arrays :
        if len ( folder ) % 2 == 0 :
            image_num = len ( folder )
            total_rows = int ( image_num / 2 )
            i = 0
            while i <= total_rows :
                doc.add_paragraph ( )
                table = doc.add_table ( rows = 0, cols = 2 )
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                row_cells = table.add_row( ).cells
                row_cells[ 0 ].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                row_cells[ 1 ].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                left_pic = row_cells[ 0 ].paragraphs[ 0 ]
                left_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_left = left_pic.add_run ( )
                run_left.add_picture ( folder[ i ], width = image_width, height = image_height )
                i += 1

                right_pic = row_cells[ 1 ].paragraphs[ 0 ]  
                right_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_right = right_pic.add_run ( )
                run_right.add_picture ( folder[ i ], width = image_width, height = image_height )

                text_row = table.add_row( ).cells
                left_cell = text_row[ 0 ].paragraphs[ 0 ]
                font = left_cell.style.font
                font.name = font_style
                font.size = font_size
                doc.add_page_break ( )
                i += 1

        elif len( folder ) % 2 != 0 :
            image_num = len ( folder )
            total_rows = int ( (image_num / 2) + 1 )
            i = 0
            while i < total_rows :
                doc.add_paragraph ( )
                table = doc.add_table ( rows = 0, cols = 2 )
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                row_cells = table.add_row( ).cells
                row_cells[ 0 ].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                row_cells[ 1 ].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                left_pic = row_cells[ 0 ].paragraphs[ 0 ]
                left_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_left = left_pic.add_run ( )
                run_left.add_picture ( folder[ i ], width = image_width, height = image_height )
                i += 1
                right_pic = row_cells[ 1 ].paragraphs[ 0 ]
                right_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_right = right_pic.add_run ( )
                run_right.add_picture ( folder[ i ], width = image_width, height = image_height )

                text_row = table.add_row( ).cells
                left_cell = text_row[ 0 ].paragraphs[ 0 ]
                font = left_cell.style.font
                font.name = font_style
                font.size = font_size
                doc.add_page_break ( )
                i += 1
                
            doc.add_paragraph ( )
            table = doc.add_table ( rows = 0, cols = 2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            row_cells = table.add_row ( ).cells
            row_cells[ 0 ].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row_cells[ 1 ].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            left_pic = row_cells[ 0 ].paragraphs[ 0 ]
            left_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_left = left_pic.add_run ( )
            run_left.add_picture ( folder[ i ], width = image_width, height = image_height )
            text_row = table.add_row( ).cells
            left_cell = text_row[ 0 ].paragraphs[ 0 ]
            font = left_cell.style.font
            font.name = font_style
            font.size = font_size
            doc.add_page_break ( )

if __name__ == "__main__" :   
    setOrientation( )             
    main( )

doc.save ( "test.docx" )